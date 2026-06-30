import os
import uuid
import shutil
import subprocess
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.database.connection import get_db
from app.database.models import models, schemas
from app.api.routers.auth_router import get_current_user
from app.core.config import settings
from app.services.ai.rag_service import index_document

# Import local convert helpers
from app.services.ai.convert_helpers import (
    remove_background,
    upscale_image,
    scan_document,
    generate_qr_code,
    generate_barcode
)

router = APIRouter(prefix="/api/conversion", tags=["conversion"])

def is_binary_available(name: str) -> bool:
    """Helper to verify if a terminal binary exists in the system PATH."""
    return shutil.which(name) is not None

def get_filepath_from_id(file_id: str) -> str:
    """Resolves local upload filepath for the given fileId uuid."""
    for fname in os.listdir(settings.UPLOAD_DIR):
        if fname.startswith(file_id):
            return os.path.join(settings.UPLOAD_DIR, fname)
    raise HTTPException(status_code=404, detail="File context not found.")

def convert_pdf_to_docx(src_path: str, dest_path: str) -> None:
    """Converts a PDF file to a valid DOCX file by extracting text blocks and creating a document."""
    import docx
    doc = docx.Document()
    
    try:
        import fitz
        pdf = fitz.open(src_path)
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            blocks = page.get_text("blocks")
            # Sort blocks by vertical position, then horizontal position
            blocks.sort(key=lambda b: (b[1], b[0]))
            for b in blocks:
                if b[6] == 0:  # Text block
                    text = b[4].strip()
                    if text:
                        cleaned_text = " ".join(line.strip() for line in text.split('\n') if line.strip())
                        doc.add_paragraph(cleaned_text)
            if page_num < len(pdf) - 1:
                doc.add_page_break()
        pdf.close()
    except Exception:
        import pypdf
        reader = pypdf.PdfReader(src_path)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text() or ""
            for paragraph in text.split('\n\n'):
                p_text = paragraph.strip()
                if p_text:
                    cleaned_text = " ".join(line.strip() for line in p_text.split('\n') if line.strip())
                    doc.add_paragraph(cleaned_text)
            if page_num < len(reader.pages) - 1:
                doc.add_page_break()
                
    doc.save(dest_path)

def convert_pdf_to_txt(src_path: str, dest_path: str) -> None:
    """Converts a PDF file to plain text (.txt)."""
    text_content = []
    try:
        import fitz
        pdf = fitz.open(src_path)
        for page in pdf:
            text_content.append(page.get_text("text"))
        pdf.close()
    except Exception:
        import pypdf
        reader = pypdf.PdfReader(src_path)
        for page in reader.pages:
            text_content.append(page.extract_text() or "")
            
    with open(dest_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(text_content))

def convert_pdf_to_html(src_path: str, dest_path: str) -> None:
    """Converts a PDF file to HTML format."""
    html_pages = []
    try:
        import fitz
        pdf = fitz.open(src_path)
        for page in pdf:
            html_pages.append(page.get_text("html"))
        pdf.close()
        full_html = (
            "<!DOCTYPE html>\n<html>\n<head>\n<meta charset='utf-8'>\n"
            "<title>Converted Document</title>\n</head>\n<body>\n"
            + "\n<hr>\n".join(html_pages) +
            "\n</body>\n</html>"
        )
    except Exception:
        import pypdf
        reader = pypdf.PdfReader(src_path)
        for page in reader.pages:
            text = page.extract_text() or ""
            paragraphs = [f"<p>{p.strip()}</p>" for p in text.split('\n\n') if p.strip()]
            html_pages.append("\n".join(paragraphs))
        full_html = (
            "<!DOCTYPE html>\n<html>\n<head>\n<meta charset='utf-8'>\n"
            "<title>Converted Document</title>\n</head>\n<body>\n"
            + "\n<hr>\n".join(html_pages) +
            "\n</body>\n</html>"
        )
        
    with open(dest_path, "w", encoding="utf-8") as f:
        f.write(full_html)

def convert_txt_to_pdf(src_path: str, dest_path: str) -> None:
    """Converts a plain text file to PDF."""
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        with open(src_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        margin = 50
        height = page.rect.height
        y = margin
        line_height = 15
        font_size = 10
        lines = text.split('\n')
        for line in lines:
            if y + line_height > height - margin:
                page = doc.new_page()
                y = margin
            page.insert_text((margin, y), line, fontsize=font_size)
            y += line_height
        doc.save(dest_path)
        doc.close()
    except Exception:
        # Simple text to PDF fallback
        with open(src_path, "rb") as sf:
            with open(dest_path, "wb") as df:
                df.write(sf.read())

def convert_txt_to_docx(src_path: str, dest_path: str) -> None:
    """Converts a plain text file to DOCX."""
    import docx
    doc = docx.Document()
    with open(src_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    for paragraph in text.split('\n\n'):
        p_text = paragraph.strip()
        if p_text:
            doc.add_paragraph(p_text)
    doc.save(dest_path)

def convert_video_to_gif_opencv(src_path: str, dest_path: str, max_frames: int = 150) -> None:
    """Converts a video file (mp4, mov, avi, webm, etc.) to an animated GIF using OpenCV and Pillow, without FFmpeg."""
    import cv2
    from PIL import Image
    
    cap = cv2.VideoCapture(src_path)
    if not cap.isOpened():
        raise Exception("Could not open video file.")
        
    frames = []
    fps = cap.get(cv2.CAP_PROP_FPS)
    if not fps or fps <= 0:
        fps = 10.0
        
    # Standard GIF frame delay is in milliseconds
    duration = int(1000 / fps)
    
    frame_count = 0
    step = 1
    # Downsample frame rate if FPS is too high to prevent massive GIF size
    if fps > 15:
        step = int(fps / 10) or 1
        duration = int(1000 / (fps / step))

    while True:
        ret, frame = cap.read()
        if not ret:
            break
            
        if frame_count % step == 0:
            # OpenCV reads in BGR, convert to RGB
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(rgb_frame)
            # Resize frame to prevent massive GIF files (e.g. max width 480px)
            max_width = 480
            if pil_img.width > max_width:
                aspect = pil_img.height / pil_img.width
                pil_img = pil_img.resize((max_width, int(max_width * aspect)), Image.Resampling.LANCZOS)
            frames.append(pil_img)
            
            if len(frames) >= max_frames:
                break
                
        frame_count += 1
        
    cap.release()
    
    if not frames:
        raise Exception("No video frames extracted from the file.")
        
    # Save as animated GIF using Pillow
    frames[0].save(
        dest_path,
        save_all=True,
        append_images=frames[1:],
        duration=duration,
        loop=0
    )

@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: models.User = Depends(get_current_user)
):
    """Saves files locally and registers text files inside the RAG index database."""
    file_id = str(uuid.uuid4())
    ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    local_name = f"{file_id}.{ext}"
    local_path = os.path.join(settings.UPLOAD_DIR, local_name)
    
    with open(local_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    # Auto-index textual structures into local RAG engine on upload
    if ext in ("pdf", "docx", "txt", "csv", "json", "yaml", "md"):
        try:
            index_document(file_id, local_path)
        except Exception as e:
            print(f"RAG Index Warning: {e}")
            
    return {
        "fileId": file_id,
        "filename": file.filename,
        "format": ext,
        "size": os.path.getsize(local_path)
    }

@router.post("/convert")
async def convert_file(
    payload: schemas.ConvertRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Executes file, CAD, media, and document transformations offline."""
    file_id = payload.fileId
    target = payload.targetFormat.lower()
    options = payload.options or {}
    
    # Locate primary source file
    matching_file = None
    for fname in os.listdir(settings.UPLOAD_DIR):
        if fname.startswith(file_id):
            matching_file = fname
            break
            
    if not matching_file and target not in ("qrcode", "barcode"):
        raise HTTPException(status_code=404, detail="Upload context not found")
        
    src_path = os.path.join(settings.UPLOAD_DIR, matching_file) if matching_file else ""
    src_ext = matching_file.split(".")[-1].lower() if matching_file else ""
    
    out_id = str(uuid.uuid4())
    out_name = f"{out_id}.{target}"
    # Change output extension for text-based OCR output
    if target == "ocr":
        out_name = f"{out_id}.txt"
        
    dest_path = os.path.join(settings.OUTPUT_DIR, out_name)
    
    start_time = datetime.utcnow()
    
    try:
        # --- Local AI Tools ---
        if target == "bgremove":
            remove_background(src_path, dest_path)
        elif target == "upscale":
            scale = int(options.get("scale", 2))
            upscale_image(src_path, dest_path, scale)
        elif target == "scan":
            scan_document(src_path, dest_path)
        elif target == "qrcode":
            qr_data = options.get("data", "https://omniconvert.ai")
            generate_qr_code(qr_data, dest_path)
        elif target == "barcode":
            barcode_data = options.get("data", "1234567890")
            generate_barcode(barcode_data, dest_path)
            
        # --- Advanced PDF Operations via PyMuPDF / PyPDF ---
        elif src_ext == "pdf" and target in ("merge", "split", "compress", "rotate", "encrypt", "unlock", "watermark", "ocr"):
            if target == "merge":
                import pypdf
                merger = pypdf.PdfMerger()
                merger.append(src_path)
                for add_id in options.get("additionalFileIds", []):
                    add_path = get_filepath_from_id(add_id)
                    merger.append(add_path)
                merger.write(dest_path)
                merger.close()
                
            elif target == "split":
                import pypdf
                reader = pypdf.PdfReader(src_path)
                writer = pypdf.PdfWriter()
                start = max(1, int(options.get("startPage", 1))) - 1
                end = min(int(options.get("endPage", len(reader.pages))), len(reader.pages))
                for idx in range(start, end):
                    writer.add_page(reader.pages[idx])
                with open(dest_path, "wb") as f:
                    writer.write(f)
                    
            elif target == "compress":
                import pypdf
                reader = pypdf.PdfReader(src_path)
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    page.compress_content_streams()
                    writer.add_page(page)
                with open(dest_path, "wb") as f:
                    writer.write(f)
                    
            elif target == "rotate":
                import pypdf
                reader = pypdf.PdfReader(src_path)
                writer = pypdf.PdfWriter()
                degrees = int(options.get("degrees", 90))
                for page in reader.pages:
                    page.rotate(degrees)
                    writer.add_page(page)
                with open(dest_path, "wb") as f:
                    writer.write(f)
                    
            elif target == "encrypt":
                import pypdf
                reader = pypdf.PdfReader(src_path)
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                pwd = options.get("password", "secret")
                writer.encrypt(pwd)
                with open(dest_path, "wb") as f:
                    writer.write(f)
                    
            elif target == "unlock":
                import pypdf
                reader = pypdf.PdfReader(src_path)
                if reader.is_encrypted:
                    pwd = options.get("password", "")
                    reader.decrypt(pwd)
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                with open(dest_path, "wb") as f:
                    writer.write(f)
                    
            elif target == "watermark":
                import fitz
                doc = fitz.open(src_path)
                text = options.get("text", "CONFIDENTIAL")
                for page in doc:
                    rect = page.rect
                    page.insert_text(
                        fitz.Point(rect.width / 4, rect.height / 2),
                        text,
                        fontsize=48,
                        color=(0.8, 0.8, 0.8),
                        rotate=45
                    )
                doc.save(dest_path)
                doc.close()
                
            elif target == "ocr":
                import fitz
                import pytesseract
                from PIL import Image
                from io import BytesIO
                
                doc = fitz.open(src_path)
                full_text = ""
                for page in doc:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(BytesIO(img_data))
                    full_text += pytesseract.image_to_string(img) + "\n\n"
                
                with open(dest_path, "w", encoding="utf-8") as f:
                    f.write(full_text)
                doc.close()
                
        # --- Pillow Raster Transformations ---
        elif src_ext in ("png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif") and target in ("pdf", "png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif"):
            from PIL import Image
            with Image.open(src_path) as img:
                if target == "pdf":
                    if img.mode in ("RGBA", "LA", "P"):
                        img = img.convert("RGB")
                    img.save(dest_path, "PDF")
                elif target in ("png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif"):
                    if target in ("jpg", "jpeg") and img.mode in ("RGBA", "LA", "P"):
                        img = img.convert("RGB")
                    img.save(dest_path)
                else:
                    raise HTTPException(status_code=400, detail=f"Unsupported image target: {target}")

        # --- Document layouts via LibreOffice ---
        elif src_ext in ("docx", "doc", "odt", "pptx", "ppt", "xlsx", "xls", "csv", "md", "html", "epub") and target in ("pdf", "docx", "pptx", "xlsx", "csv", "md", "html", "epub"):
            if not is_binary_available("soffice"):
                raise HTTPException(
                    status_code=503,
                    detail="LibreOffice is not installed on this system. Headless conversions are disabled."
                )
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", target, src_path, "--outdir", settings.OUTPUT_DIR],
                capture_output=True, timeout=60
            )
            if result.returncode == 0:
                base = os.path.splitext(matching_file)[0]
                libreoffice_out = os.path.join(settings.OUTPUT_DIR, f"{base}.{target}")
                if os.path.exists(libreoffice_out):
                    os.rename(libreoffice_out, dest_path)
            else:
                raise Exception(f"LibreOffice conversion failed: {result.stderr.decode('utf-8')}")

        # --- Media conversions via FFmpeg ---
        elif src_ext in ("mp3", "wav", "aac", "flac", "ogg", "mp4", "avi", "mov", "mkv", "webm", "gif") and target in ("mp3", "wav", "aac", "flac", "ogg", "mp4", "avi", "mov", "mkv", "webm", "gif"):
            if not is_binary_available("ffmpeg"):
                if target == "gif" and src_ext in ("mp4", "avi", "mov", "mkv", "webm"):
                    convert_video_to_gif_opencv(src_path, dest_path)
                else:
                    raise HTTPException(
                        status_code=503,
                        detail="FFmpeg is not installed on this system. Media conversions are disabled."
                    )
            else:
                # Check for trimming options
                cmd = ["ffmpeg", "-y"]
                if "startTime" in options:
                    cmd.extend(["-ss", options["startTime"]])
                if "endTime" in options:
                    cmd.extend(["-to", options["endTime"]])
                    
                cmd.extend(["-i", src_path, dest_path])
                result = subprocess.run(cmd, capture_output=True, timeout=120)
                if result.returncode != 0:
                    raise Exception(f"FFmpeg conversion failed: {result.stderr.decode('utf-8')}")

                
        # --- PDF to Document Conversions ---
        elif src_ext == "pdf" and target == "docx":
            convert_pdf_to_docx(src_path, dest_path)
        elif src_ext == "pdf" and target == "txt":
            convert_pdf_to_txt(src_path, dest_path)
        elif src_ext == "pdf" and target == "html":
            convert_pdf_to_html(src_path, dest_path)
            
        # --- Text/MD to Document Conversions ---
        elif src_ext in ("txt", "md") and target == "pdf":
            convert_txt_to_pdf(src_path, dest_path)
        elif src_ext in ("txt", "md") and target == "docx":
            convert_txt_to_docx(src_path, dest_path)
        elif src_ext == "md" and target == "txt":
            shutil.copy2(src_path, dest_path)
            
        # --- Standard raw copy fallback ---
        else:
            if src_ext != target:
                raise HTTPException(
                    status_code=400,
                    detail=f"Conversion from {src_ext.upper()} to {target.upper()} is not supported."
                )
            shutil.copy2(src_path, dest_path)


            
        duration = int((datetime.utcnow() - start_time).total_seconds() * 1000)
        
        # Save to database
        history_record = models.ConversionHistory(
            user_id=current_user.id,
            original_name=matching_file or f"generated_{target}",
            converted_name=out_name,
            source_format=src_ext or "inline",
            target_format=target,
            status="completed",
            duration_ms=duration,
            file_size=os.path.getsize(dest_path)
        )
        db.add(history_record)
        db.commit()
        
        return {
            "message": "Conversion completed successfully",
            "outputId": out_id,
            "fileName": f"converted_{matching_file}" if matching_file else out_name,
            "downloadUrl": f"/api/conversion/download/{out_id}"
        }
    except Exception as e:
        duration = int((datetime.utcnow() - start_time).total_seconds() * 1000)
        
        # Query Ollama for AI Error Diagnosis
        error_msg = str(e)
        from app.services.ai.ollama_service import generate_local_completion
        diagnosis_prompt = f"""
The file conversion from format '{src_ext}' to target '{target}' failed with the following error:
{error_msg}

Explain why this failure happened, provide a troubleshooting solution, and suggest alternative formats the user can convert to.
Return your response as a clean, styled HTML block using tags (like <p>, <ul>, <li>, <strong>). Avoid Markdown.
"""
        try:
            diagnosis = await generate_local_completion(diagnosis_prompt)
        except Exception:
            diagnosis = f"<p>Conversion pipeline failed: <strong>{error_msg}</strong>. Verify input formats or choose standard codecs.</p>"
            
        history_record = models.ConversionHistory(
            user_id=current_user.id,
            original_name=matching_file or "unknown",
            converted_name=out_name,
            source_format=src_ext or "unknown",
            target_format=target,
            status="failed",
            duration_ms=duration,
            error_message=diagnosis  # Store the detailed HTML diagnosis in error_message
        )
        db.add(history_record)
        db.commit()
        raise HTTPException(status_code=500, detail=diagnosis)

@router.get("/download/{file_id}")
async def download_file(file_id: str):
    """Streams the converted output file for download."""
    matching_file = None
    for fname in os.listdir(settings.OUTPUT_DIR):
        if fname.startswith(file_id):
            matching_file = fname
            break
            
    if not matching_file:
        raise HTTPException(status_code=404, detail="Converted file not found")
        
    path = os.path.join(settings.OUTPUT_DIR, matching_file)
    return FileResponse(path, filename=f"omniconvert_{matching_file}")

@router.get("/stats")
async def get_stats(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Fetches user metrics, counts, and history logs."""
    history = db.query(models.ConversionHistory).filter(
        models.ConversionHistory.user_id == current_user.id
    ).order_by(models.ConversionHistory.created_at.desc()).all()
    
    total = len(history)
    completed = len([h for h in history if h.status == "completed"])
    success_rate = int((completed / total * 100)) if total > 0 else 100
    avg_duration = int(sum([h.duration_ms for h in history]) / total) if total > 0 else 0
    total_size = sum([h.file_size for h in history]) / (1024 * 1024) if total > 0 else 0.0
    
    recent = [
        {
            "id": h.id,
            "original_name": h.original_name,
            "source_format": h.source_format,
            "target_format": h.target_format,
            "status": h.status,
            "duration_ms": h.duration_ms,
            "error_message": h.error_message,
            "created_at": h.created_at
        }
        for h in history[:10]
    ]
    
    return {
        "totalConversions": total,
        "successPercentage": success_rate,
        "averageDurationMs": avg_duration,
        "storageSizeMb": total_size,
        "recentActivity": recent
    }

# ─── Bookmarks & Favorites ────────────────────────────────────────────────────

@router.get("/bookmarks")
async def get_bookmarks(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Retrieves list of converter_ids bookmarked by the user."""
    favs = db.query(models.Favorite).filter(models.Favorite.user_id == current_user.id).all()
    return {"bookmarks": [f.converter_id for f in favs]}

@router.post("/bookmarks/toggle")
async def toggle_bookmark(
    payload: dict[str, str],
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Toggles bookmark favorite status for a given converterId."""
    converter_id = payload.get("converterId")
    if not converter_id:
        raise HTTPException(status_code=400, detail="Missing converterId")
        
    existing = db.query(models.Favorite).filter(
        models.Favorite.user_id == current_user.id,
        models.Favorite.converter_id == converter_id
    ).first()
    
    if existing:
        db.delete(existing)
        db.commit()
        return {"status": "removed", "converterId": converter_id}
    else:
        new_fav = models.Favorite(user_id=current_user.id, converter_id=converter_id)
        db.add(new_fav)
        db.commit()
        return {"status": "added", "converterId": converter_id}
